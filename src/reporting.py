from subprocess import  Popen
import os
import platform
from docx import Document
from docx.shared import Inches
import base64
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx2pdf import convert
from docx.enum.section import WD_ORIENT #added
from docx.enum.section import WD_SECTION # #added
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT #added
# from docx import Document
from datetime import datetime
from pdfrw import PdfReader, PdfWriter
from docx.shared import RGBColor
from PIL import Image, ImageDraw, ImageFont
import src.report_cover as report_cover
import tempfile
import pythoncom

current_system = platform.system()

def linux_convert_to_pdf(input_docx, out_folder):
    LIBRE_OFFICE = "libreoffice"
    p = Popen(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir',
               out_folder, input_docx])
    print([LIBRE_OFFICE, '--convert-to', 'pdf', input_docx])
    p.communicate()



def convert_word_to_pdf(word_path, pdf_path):
    pythoncom.CoInitialize()
    try:
        convert(word_path, pdf_path)
    finally:
        pythoncom.CoUninitialize()
       
       
def save_visualization_to_file(visualization_base64, index):
   image_data = base64.b64decode(visualization_base64)
   image_path = f'visualization{index}.png'
   with open(image_path, 'wb') as file:
       file.write(image_data)
   return image_path


def add_border_to_table_cell(cell, border_color="#808080", border_width="5"):
   # This function adds a border to a table cell
   sides = ['top', 'left', 'bottom', 'right']
   for side in sides:
       tag = 'w:{}'.format(side)
       border_elm = OxmlElement(tag)
       border_elm.set(qn('w:val'), 'single')
       border_elm.set(qn('w:sz'), border_width)
       border_elm.set(qn('w:color'), border_color)
       border_elm.set(qn('w:space'), "0")
       cell._tc.get_or_add_tcPr().append(border_elm)          
       
def _removed_without_path_records(visualizations):
    blank_vis = []
    for index in range(len(visualizations)):
     record = visualizations[index]
     if "visualizations_path" not in record:
          blank_vis.append(index)
          
    for blank_index in blank_vis:
        del visualizations[blank_index]
    return visualizations


def add_visualizations_to_doc(document, visualizations):
    kpi_heading = document.add_paragraph()
    kpi_heading.add_run("KPI Report:").bold = True
    for run in kpi_heading.runs:
       run.font.color.rgb = RGBColor(0, 0, 255)  # Set color to blue
       run.font.size = Pt(15)
    max_len=len(visualizations)-1
    cnt=0
    for index in range(len(visualizations)):
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
            img = Image.open(visualizations[index]["visualizations_path"])
            img.save(tmp_file.name)        
            document.add_picture(tmp_file.name, width=Inches(6.0))
            # if "visualizations_description" in visualizations[index]:
            #     table = document.add_table(rows=2, cols=3)
            #     col_count = 0
            #     for desctiption in visualizations[index]["visualizations_description"]:
            #         cell = table.cell(0, col_count)
            #         cell.margin_top = Inches(0.1)
            #         cell.margin_bottom = Inches(0.1)
            #         cell.margin_left = Inches(0.1)
            #         cell.margin_right = Inches(0.1)
            #         add_border_to_table_cell(cell, "#808080", "5")  # Set border color and width
            #         paragraph = cell.paragraphs[0]
            #         run = paragraph.add_run(desctiption["section"])

            #         cell = table.cell(1, col_count)
            #         cell.margin_top = Inches(0.1)
            #         cell.margin_bottom = Inches(0.1)
            #         cell.margin_left = Inches(0.1)
            #         cell.margin_right = Inches(0.1)
            #         add_border_to_table_cell(cell, "#808080", "5")  # Set border color and width
            #         paragraph = cell.paragraphs[0]
            #         run = paragraph.add_run(desctiption["explanation"])

            #         col_count += 1
            
            if cnt<(max_len):
                document.add_paragraph("________________________________________________________________________________________________________")
            cnt+=1

 
def add_selections_summary_kpis(doc, tenant_name, applications, start_date, end_date):
   p = doc.add_paragraph()
   run = p.add_run(f'Tenant Name: {tenant_name}')
   run.bold = True
   p = doc.add_paragraph()
   run = p.add_run(f'Selected Applications: {", ".join(applications)}')
   run.bold = True
   p = doc.add_paragraph()
   run = p.add_run(f'Start Date: {start_date}')
   run.bold = True
   p = doc.add_paragraph()
   run = p.add_run(f'End Date: {end_date}')
   run.bold = True
   # Space between end_date and summary
   doc.add_paragraph(' ')
#    summary_para = doc.add_paragraph()
#    summary_para.add_run("Summary:").bold = True
#    for run in summary_para.runs:
#        run.font.color.rgb = RGBColor(0, 0, 255)  # Set color to blue
#        run.font.size = Pt(15)
#    doc.add_paragraph(textual_summary)
#    doc.add_paragraph("")
   # Space between summary and KPIs

def date_range(start_date, end_date):
    

    # Extract month and year from start and end dates
    start_month_year = start_date.strftime("%B %Y")
    end_month_year = end_date.strftime("%B %Y")

    # Return date range string
    return f"{start_month_year} to {end_month_year}"

# write given list of text into image and returns image path
def write_text_center(image_pth, text): 
    print("recieved img path:",image_pth)
    font = ImageFont.load_default(50)
    W= int(2480/2.92)
    H= int(3508/3.2)

    im = Image.open(image_pth)
    #Image.open(image_pth)
    
    resized_image = im.resize((W, H))
    draw = ImageDraw.Draw(resized_image)
    _, _, w, h = draw.textbbox((0, 30), text[0], font=font)
    draw.text(((W-w)/2, (H-h)/2), text[0], font=font, fill='black')
    
    font = ImageFont.load_default(30)
    _, _, w, h = draw.textbbox((0, -60), text[1], font=font)
    draw.text(((W-w)/2, (H-h)/2), text[1], font=font, fill='black')
    front_img_path="./front_img2.png"
    resized_image.save(front_img_path)

    return front_img_path

#Makes reports/documents first page and sets some parameters of remaining pages returns "docx" object
def make_document_first_page(doc,image_path,start_date,end_date,Message):

    range_date=date_range(start_date,end_date)
    myMessage = [Message,range_date]
    image_path=write_text_center(image_path,myMessage)

    sec=doc.add_section()# this page is required and intentionally left blank.
    new_section = doc.sections[1]
    new_section.left_margin=Inches(0)
    new_section.right_margin=Inches(0)
    new_section.top_margin=Inches(0)
    new_section.bottom_margin=Inches(0)
    doc.add_picture(image_path, width=new_section.page_width, height=10089400)
    doc.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # delete image
    if os.path.isfile(image_path):
        os.remove(image_path)
    
    
    # remaining page start
    new_section=doc.add_section()
    new_section.left_margin=Inches(0.3)
    new_section.right_margin=Inches(0.3)
    new_section.top_margin=Inches(0.3)
    new_section.bottom_margin=Inches(0.1)  

    new_section.start_type = WD_SECTION.CONTINUOUS
    header = new_section.header
    header.is_linked_to_previous = False
    header_paragraph = header.add_paragraph()
    
    # Add the header text to the paragraph
    run1 = header_paragraph.add_run("\t\t\t\t\tPartsOps Dashboard")
    run1.font.color.rgb = RGBColor(50,200,10)
    run1.font.size = Pt(25)
    run1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    section = doc.sections[-1]
    # Create a footer
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Add the current date and time to the footer
    current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    footer_paragraph.text = f"Report generated on : {current_time}"

    return doc   

#Removes first page of pdf and returns modified pdf's path
def remove_blank_page(pdf_path):
    # Define the reader and writer objects
    reader_input = PdfReader(pdf_path)
    writer_output = PdfWriter()

    # Go through the pages one after the next
    for current_page in range(len(reader_input.pages)):
        if current_page != 0:# Removing "0" th page
            writer_output.addpage(reader_input.pages[current_page])
    # Write the modified content to disk
    writer_output.write(pdf_path)

    return pdf_path

 
def generate_report(tenant_name, applications, start_date, end_date, vis_details):
   doc = Document()
   image_path = "images/cover_pages/base_image_1.jpg"
   Message = "CIDDS"
   doc=make_document_first_page(doc,image_path,start_date,end_date,Message)

#    doc.add_paragraph() 
#    header = doc.sections[0].header
#    header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
#    header_paragraph.text = "PartsOps Dashboard"
#    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
#    for run in header_paragraph.runs:
#        run.font.color.rgb = RGBColor(0, 0, 255)  # Set color to blue
#        run.font.size = Pt(20)

#    section = doc.sections[-1]
#    # Create a footer
#    footer = section.footer
#    footer_paragraph = footer.paragraphs[0]
#    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

#    # Add the current date and time to the footer
#    current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
#    footer_paragraph.text = f"Report generated on : {current_time}"
 
   
#    add_selections_summary_kpis(doc, tenant_name, applications, start_date, end_date)
   vis_details = _removed_without_path_records(vis_details)
   add_visualizations_to_doc(doc,vis_details)
   file_name = f"KPI_Report_{datetime.now():%Y-%m-%d_%H-%M-%S}"
   doc_name = f'{file_name}.docx'
   doc_path = f"data/reports/{doc_name}"
   pdf_folder_path = f'data/reports'
   doc.save(doc_path)
   pdf_path = f"{pdf_folder_path}/{file_name}.pdf"

   # pdf_path = os.path.join(os.path.dirname(doc_path), pdf_filename)
   if current_system == 'Linux':
       linux_convert_to_pdf(doc_path, pdf_folder_path)
   if current_system == 'Windows':
       convert_word_to_pdf(doc_path, pdf_path)
   
   pdf_path = remove_blank_page(pdf_path)
   print(pdf_path)
       # Open the PDF file and create a download button
#    pdf_path = report_cover.generate_cover(start_date, end_date, tenant_name, pdf_path)
   return pdf_path

if __name__ == "__main__":
    tenant_name = 'test-tenant_name'
    applications = ['test-application']
    start_date = 'test-sd'
    end_date = 'test-ed'
  
    kpi_texts = ['kpi-text1','kpi-text2']
    visualization_paths= ['data/graphs/visualization_0.png','data/graphs/visualization_0.png']
    generate_report(tenant_name, applications, start_date, end_date,kpi_texts,visualization_paths)
    
    
  
  
             
                      




   

                  
                  
       